import sys
try:
    import docx
except ImportError:
    pass

try:
    doc = docx.Document(r"d:\Avengers Doomsday\Assignment-6\Borole Druv.docx")
    print("--- WORD PARAGRAPHS ---")
    for p in doc.paragraphs:
        if p.text.strip():
            print(repr(p.text))
            
    print("--- WORD TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            print([c.text.strip() for c in row.cells])
except Exception as e:
    print("Error:", e)
