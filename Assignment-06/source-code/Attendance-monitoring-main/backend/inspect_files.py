import sys
import pandas as pd
try:
    import docx
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "openpyxl"])
    import docx

print("---- EXCEL ----")
try:
    df = pd.read_excel(r"d:\Avengers Doomsday\Assignment-6\Attendance_Report_37.xlsx", header=None)
    for i in range(min(50, len(df))):
        print(df.iloc[i].values)
except Exception as e:
    print("Error reading excel:", e)

print("---- WORD ----")
try:
    doc = docx.Document(r"d:\Avengers Doomsday\Assignment-6\Borole Druv.docx")
    for para in doc.paragraphs:
        print(para.text)
    for table in doc.tables:
        print("--- Table ---")
        for row in table.rows:
            print([cell.text for cell in row.cells])
except Exception as e:
    print("Error reading docx:", e)
