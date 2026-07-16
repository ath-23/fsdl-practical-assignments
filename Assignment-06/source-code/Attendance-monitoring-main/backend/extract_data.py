import sys
import pandas as pd
import json

try:
    df = pd.read_excel(r"d:\Avengers Doomsday\Assignment-6\Attendance_Report_37.xlsx", header=None)
    data = {"excel_rows": df.head(50).fillna("").values.tolist()}
    with open(r"d:\Avengers Doomsday\Assignment-6\backend\excel_data.json", "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)
except Exception as e:
    with open(r"d:\Avengers Doomsday\Assignment-6\backend\excel_data.json", "w", encoding="utf-8") as f:
        f.write(json.dumps({"error": str(e)}))

try:
    import docx
    doc = docx.Document(r"d:\Avengers Doomsday\Assignment-6\Borole Druv.docx")
    word_data = {"paragraphs": [], "tables": []}
    for p in doc.paragraphs:
        if p.text.strip():
            word_data["paragraphs"].append(p.text)
    for table in doc.tables:
        t_data = []
        for row in table.rows:
            t_data.append([c.text.strip() for c in row.cells])
        word_data["tables"].append(t_data)
    with open(r"d:\Avengers Doomsday\Assignment-6\backend\word_data.json", "w", encoding="utf-8") as f:
        json.dump(word_data, f, indent=2)
except Exception as e:
    with open(r"d:\Avengers Doomsday\Assignment-6\backend\word_data.json", "w", encoding="utf-8") as f:
        f.write(json.dumps({"error": str(e)}))
